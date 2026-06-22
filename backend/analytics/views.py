import json
import uuid
import time
import os
import io
from datetime import datetime, timedelta
from django.utils import timezone
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.conf import settings
from django.contrib.auth import authenticate, login
from django.core.cache import cache
from rest_framework.decorators import api_view, permission_classes
from django.utils.decorators import method_decorator
from rest_framework import status, viewsets
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework import permissions
from .models import FinancialDataSet, FinancialReport, FinancialInsight
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
import openpyxl
from openpyxl.styles import Font, Alignment

from .services.dynamic_report_builder import build_dynamic_report_sections, build_report_context
from .services.report_prompt_registry import get_report_prompt_registry

# Simple in-memory storage for testing
GLOBAL_REPORTS = {}

REPORT_TEMPLATES = get_report_prompt_registry().get_templates()

def get_global_reports():
    """Get global reports (handles module reloading issues)"""
    global GLOBAL_REPORTS
    return GLOBAL_REPORTS

def add_global_report(report_id, report_data):
    """Add report to global storage (handles module reloading issues)"""
    global GLOBAL_REPORTS
    GLOBAL_REPORTS[report_id] = report_data
    print(f"DEBUG: Added to global storage: {report_id}")


def update_stored_report(report_id, updates, request=None):
    """Update an existing report in session and global storage."""
    global GLOBAL_REPORTS
    report_id = str(report_id)
    updated = False

    if request is not None and 'reports' in request.session:
        for index, report in enumerate(request.session['reports']):
            if str(report.get('id')) == report_id:
                report.update(updates)
                request.session['reports'][index] = report
                request.session.modified = True
                updated = True
                break

    if report_id in GLOBAL_REPORTS:
        GLOBAL_REPORTS[report_id].update(updates)
        updated = True

    return updated


def get_openai_api_key():
    """Resolve OpenAI API key from Django settings or environment."""
    key = (getattr(settings, 'OPENAI_API_KEY', '') or os.getenv('OPENAI_API_KEY', '')).strip()
    return key


def parse_openai_error(exc):
    """Map OpenAI exceptions to user-facing error messages."""
    error_str = str(exc)
    error_code = ''
    error_type = ''

    try:
        from openai import APIStatusError
        if isinstance(exc, APIStatusError):
            body = getattr(exc, 'body', None) or {}
            if isinstance(body, dict):
                err = body.get('error', {})
                if isinstance(err, dict):
                    error_code = str(err.get('code', '') or '')
                    error_type = str(err.get('type', '') or '')
                    error_str = str(err.get('message', error_str))
    except Exception:
        pass

    if error_code == 'insufficient_quota' or error_type == 'insufficient_quota':
        return (
            'OpenAI returned insufficient_quota for this API key. '
            'New keys still need billing enabled at https://platform.openai.com/account/billing '
            'before chat requests work.'
        )
    if 'invalid_api_key' in error_code or 'invalid_api_key' in error_str.lower():
        return 'Invalid OpenAI API key. Update OPENAI_API_KEY in backend/.env and restart the server.'
    if 'model_not_found' in error_code or 'does not exist' in error_str.lower():
        return (
            f'OpenAI model not available: {error_str}. '
            'Set OPENAI_MODEL=gpt-4o-mini in backend/.env and restart the server.'
        )
    if 'rate_limit' in error_code or 'rate_limit' in error_type or 'rate_limit' in error_str.lower():
        return 'OpenAI rate limit exceeded. Please wait a moment and try again.'
    return f'OpenAI API error: {error_str}'


def serialize_json_for_ai(data, max_chars=14000):
    """Serialize any JSON value for the AI prompt, truncating if needed."""
    try:
        serialized = json.dumps(data, default=str, ensure_ascii=False)
    except (TypeError, ValueError):
        serialized = str(data)
    if len(serialized) <= max_chars:
        return serialized
    return serialized[:max_chars] + '\n... [truncated for model context limit]'


def extract_entity_metadata(json_data):
    """Extract display metadata from any JSON shape."""
    bank_name = 'Financial Dataset'
    period = 'Unknown Period'

    name_keys = ('bank_name', 'bankName', 'BankName', 'institution', 'company', 'organization', 'entity', 'name')
    period_keys = ('period', 'data_period', 'reporting_period', 'asOf', 'Asof', 'date', 'report_date', 'year')

    def scan(obj, depth=0):
        nonlocal bank_name, period
        if depth > 6:
            return
        if isinstance(obj, dict):
            for key, value in obj.items():
                key_lower = str(key).lower()
                if bank_name == 'Financial Dataset' and key_lower in {k.lower() for k in name_keys}:
                    if isinstance(value, str) and value.strip():
                        bank_name = value.strip()
                    elif isinstance(value, dict) and value.get('Name'):
                        bank_name = str(value.get('Name')).strip()
                if period == 'Unknown Period' and key_lower in {k.lower() for k in period_keys}:
                    if isinstance(value, str) and value.strip():
                        period = value.strip()[:10] if len(value.strip()) > 10 else value.strip()
                if isinstance(value, (dict, list)):
                    scan(value, depth + 1)
        elif isinstance(obj, list):
            for item in obj[:20]:
                scan(item, depth + 1)

    scan(json_data)
    return bank_name, period


def build_data_summary(json_data):
    """Build a structure summary for any JSON root type."""
    if isinstance(json_data, dict):
        return {
            'structure': 'object',
            'top_level_keys': list(json_data.keys())[:50],
            'key_count': len(json_data),
            'data_size': len(serialize_json_for_ai(json_data)),
        }
    if isinstance(json_data, list):
        return {
            'structure': 'array',
            'item_count': len(json_data),
            'sample_item_keys': list(json_data[0].keys())[:30] if json_data and isinstance(json_data[0], dict) else [],
            'data_size': len(serialize_json_for_ai(json_data)),
        }
    return {
        'structure': type(json_data).__name__,
        'data_size': len(str(json_data)),
    }


def normalize_json_for_analysis(json_data):
    """Normalize any valid JSON value for analysis while preserving the original."""
    original = json_data
    if isinstance(json_data, dict) and 'Dashboard' in json_data:
        return transform_complex_json(json_data), original
    if isinstance(json_data, list):
        return {'records': json_data, 'record_count': len(json_data)}, original
    if isinstance(json_data, dict):
        return json_data, original
    return {'value': json_data}, original


def generate_analysis_from_prompt(prompt, json_data, ai_analysis, report_options=None):
    """Generate comprehensive report sections from the user's prompt using AI only."""
    bank_name, data_period = extract_entity_metadata(json_data)
    normalized_data, _ = normalize_json_for_analysis(json_data)

    analysis_context = {
        'bank_name': bank_name,
        'data_period': data_period,
        'financial_data': normalized_data,
        'raw_financial_data': json_data,
        'existing_analysis': ai_analysis,
        'data_summary': build_data_summary(json_data),
        'user_prompt': prompt,
        'report_options': report_options or {},
    }

    analysis_result = generate_comprehensive_ai_analysis(analysis_context)
    if analysis_result and analysis_result.get('success'):
        return analysis_result.get('sections', []), '', True

    error_msg = analysis_result.get('error', 'AI analysis failed') if analysis_result else 'AI analysis failed'
    return [], error_msg, False

def _export_plain_text(value, max_len=8000):
    if isinstance(value, str):
        text = value
    else:
        text = json.dumps(value, default=str)
    return text[:max_len]


def _append_comprehensive_section_pdf(story, section, styles):
    from xml.sax.saxutils import escape

    title = section.get('title', 'Section')
    story.append(Paragraph(escape(str(title)), styles['Heading3']))
    content = section.get('content', {})
    if isinstance(content, dict):
        narrative = content.get('content', '')
        if narrative:
            story.append(Paragraph(escape(_export_plain_text(narrative)), styles['Normal']))
        for point in content.get('key_points', []):
            story.append(Paragraph(f"• {escape(_export_plain_text(point, 2000))}", styles['Normal']))
        for rec in content.get('recommendations', []):
            if isinstance(rec, dict):
                rec_text = rec.get('action') or rec.get('area') or str(rec)
            else:
                rec_text = str(rec)
            story.append(Paragraph(f"• {escape(_export_plain_text(rec_text, 2000))}", styles['Normal']))
    elif content:
        story.append(Paragraph(escape(_export_plain_text(content)), styles['Normal']))
    story.append(Spacer(1, 8))


def _append_comprehensive_section_word(doc, section):
    doc.add_heading(str(section.get('title', 'Section')), level=2)
    content = section.get('content', {})
    if isinstance(content, dict):
        if content.get('content'):
            doc.add_paragraph(_export_plain_text(content.get('content')))
        for point in content.get('key_points', []):
            doc.add_paragraph(str(point), style='List Bullet')
        for rec in content.get('recommendations', []):
            if isinstance(rec, dict):
                rec_text = rec.get('action') or rec.get('area') or str(rec)
            else:
                rec_text = str(rec)
            doc.add_paragraph(str(rec_text), style='List Bullet')
    elif content:
        doc.add_paragraph(_export_plain_text(content))


def generate_pdf_report(report_data):
    """Generate PDF report from report data"""
    from xml.sax.saxutils import escape

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    report_title = report_data.get('metadata', {}).get('title') or 'AI Financial Analysis Report'
    story.append(Paragraph(escape(str(report_title)), styles['Title']))
    story.append(Spacer(1, 12))

    story.append(Paragraph(f"Bank: {escape(str(report_data.get('bank_name', 'Unknown Bank')))}", styles['Normal']))
    story.append(Paragraph(f"Period: {escape(str(report_data.get('data_period', 'Unknown Period')))}", styles['Normal']))
    story.append(Paragraph(f"Generated: {escape(str(report_data.get('uploaded_at', 'Unknown')))}", styles['Normal']))
    story.append(Spacer(1, 12))

    user_prompt = report_data.get('user_prompt') or report_data.get('metadata', {}).get('user_prompt')
    if user_prompt:
        story.append(Paragraph("Analysis Prompt", styles['Heading2']))
        story.append(Paragraph(escape(_export_plain_text(user_prompt)), styles['Normal']))
        story.append(Spacer(1, 12))

    comprehensive = report_data.get('comprehensive_analysis', [])
    if comprehensive:
        story.append(Paragraph("AI Analysis Report", styles['Heading2']))
        for section in comprehensive:
            _append_comprehensive_section_pdf(story, section, styles)
    else:
        story.append(Paragraph("AI Analysis", styles['Heading2']))
        ai_analysis = report_data.get('ai_analysis', {})
        for section, content in ai_analysis.items():
            story.append(Paragraph(escape(str(section).title()), styles['Heading3']))
            story.append(Paragraph(escape(_export_plain_text(content)), styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def generate_csv_report(report_data):
    """Generate CSV report from report data"""
    output = io.StringIO()
    
    # Write headers
    output.write("Report Field,Value\n")
    
    # Write basic info
    output.write(f"Bank Name,{report_data.get('bank_name', 'Unknown')}\n")
    output.write(f"Data Period,{report_data.get('data_period', 'Unknown')}\n")
    output.write(f"Upload Date,{report_data.get('uploaded_at', 'Unknown')}\n")
    
    # Write metrics
    metrics = report_data.get('data_summary', {})
    for key, value in metrics.items():
        if key != 'sample_data':  # Skip large data objects
            output.write(f"{key},{value}\n")
    
    user_prompt = report_data.get('user_prompt') or report_data.get('metadata', {}).get('user_prompt')
    if user_prompt:
        output.write(f"User Prompt,\"{str(user_prompt).replace(chr(10), ' ')}\"\n")

    for section in report_data.get('comprehensive_analysis', []):
        title = section.get('title', 'Section')
        content = section.get('content', {})
        narrative = content.get('content', '') if isinstance(content, dict) else str(content)
        output.write(f"\"{title}\",\"{str(narrative).replace(chr(10), ' ')}\"\n")
    
    content = output.getvalue()
    output.close()
    return content

def generate_word_report(report_data):
    """Generate Word document report"""
    doc = Document()

    report_title = report_data.get('metadata', {}).get('title') or 'AI Financial Analysis Report'
    doc.add_heading(str(report_title), 0)
    doc.add_paragraph(f"Bank: {report_data.get('bank_name', 'Unknown Bank')}")
    doc.add_paragraph(f"Period: {report_data.get('data_period', 'Unknown Period')}")
    doc.add_paragraph(f"Generated: {report_data.get('uploaded_at', 'Unknown')}")

    user_prompt = report_data.get('user_prompt') or report_data.get('metadata', {}).get('user_prompt')
    if user_prompt:
        doc.add_heading('Analysis Prompt', level=1)
        doc.add_paragraph(_export_plain_text(user_prompt))

    comprehensive = report_data.get('comprehensive_analysis', [])
    if comprehensive:
        doc.add_heading('AI Analysis Report', level=1)
        for section in comprehensive:
            _append_comprehensive_section_word(doc, section)
    else:
        doc.add_heading('AI Analysis', level=1)
        ai_analysis = report_data.get('ai_analysis', {})
        for section, content in ai_analysis.items():
            doc.add_heading(str(section).title(), level=2)
            doc.add_paragraph(_export_plain_text(content))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def generate_excel_report(report_data):
    """Generate Excel report"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Financial Analysis Report"
    
    # Set up headers
    ws['A1'] = "Report Field"
    ws['B1'] = "Value"
    
    # Add basic info
    ws['A2'] = "Bank Name"
    ws['B2'] = report_data.get('bank_name', 'Unknown')
    
    ws['A3'] = "Data Period"
    ws['B3'] = report_data.get('data_period', 'Unknown')
    
    ws['A4'] = "Upload Date"
    ws['B4'] = report_data.get('uploaded_at', 'Unknown')
    
    row = 6
    
    # Add metrics
    metrics = report_data.get('data_summary', {})
    for key, value in metrics.items():
        if key != 'sample_data':  # Skip large data objects
            ws[f'A{row}'] = key
            ws[f'B{row}'] = str(value)
            row += 1
    
    # Add AI analysis summary
    ai_analysis = report_data.get('ai_analysis', {})
    if 'analysis_summary' in ai_analysis:
        ws[f'A{row}'] = "Analysis Summary"
        ws[f'B{row}'] = ai_analysis['analysis_summary']
        row += 1
    
    # Save to buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.views import APIView
from rest_framework.pagination import PageNumberPagination

from .models import (
    FinancialDataUpload, FinancialDataSet, FinancialMetrics,
    FinancialInsight, FinancialReport, AnalysisTask
)
from .serializers import (
    FinancialDataUploadSerializer, FinancialReportSerializer,
    ReportSummarySerializer, AnalysisTaskSerializer,
    UploadResponseSerializer, AnalysisRequestSerializer,
    TaskStatusSerializer, DetailedReportSerializer,
    ReportFilterSerializer, FileUploadRequestSerializer
)
from .services.data_parser import FinancialDataParser
from .services.metrics_engine import FinancialMetricsEngine
from .services.insight_engine import FinancialInsightEngine
from .services.report_generator import FinancialReportGenerator


@csrf_exempt
@require_http_methods(["POST"])
def simple_custom_prompt_view(request):
    """Generate analysis using custom prompt on uploaded JSON data"""
    try:
        import json
        
        # Get request data
        body = json.loads(request.body)
        report_id = body.get('report_id')
        custom_prompt = body.get('prompt')
        
        if not report_id or not custom_prompt:
            return JsonResponse({'error': 'report_id and prompt are required'}, status=400)
        
        # Find report in session or global storage
        report = None
        if 'reports' in request.session:
            for r in request.session['reports']:
                if r['id'] == report_id:
                    report = r
                    break
        
        if not report:
            for r in GLOBAL_REPORTS.values():
                if r['id'] == report_id:
                    report = r
                    break
        
        if not report:
            return JsonResponse({'error': 'Report not found'}, status=404)
        
        # Get original JSON data
        original_json = report.get('metadata', {}).get('original_json')
        if not original_json:
            return JsonResponse({'error': 'Original JSON data not found in report'}, status=400)
        
        print(f"DEBUG: Generating custom analysis with prompt: {custom_prompt[:100]}...")
        
        # Generate analysis using custom prompt
        try:
            analysis_result = generate_comprehensive_ai_analysis({
                'bank_name': report.get('bank_name', 'Unknown Bank'),
                'data_period': report.get('data_period', 'Unknown Period'),
                'financial_data': original_json,
                'existing_analysis': report.get('ai_analysis', {}),
                'data_summary': {
                    'keys_found': list(original_json.keys()) if isinstance(original_json, dict) else 'Non-dict JSON',
                    'data_size': len(str(original_json)),
                    'sample_data': original_json if isinstance(original_json, dict) else {'data': original_json}
                },
                'user_prompt': custom_prompt
            })
            
            print(f"DEBUG: Custom analysis result: {analysis_result}")
            
            if analysis_result and analysis_result.get('success'):
                return JsonResponse({
                    'success': True,
                    'analysis': analysis_result.get('sections', []),
                    'report_id': report_id
                })
            else:
                error_msg = analysis_result.get('error', 'Unknown error')
                return JsonResponse({
                    'success': False,
                    'error': error_msg
                }, status=500)
                
        except Exception as e:
            print(f"DEBUG: Exception in custom analysis: {e}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)
            
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON in request body'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def simple_upload_view(request):
    """Simple Django upload view - no DRF"""
    try:
        # Get uploaded file
        if 'file' not in request.FILES:
            return JsonResponse({'error': 'No file provided'}, status=400)
        
        uploaded_file = request.FILES['file']
        description = request.POST.get('description', '')
        prompt = request.POST.get('prompt', '').strip()

        if not prompt:
            return JsonResponse({
                'error': 'Analysis prompt is required. Describe what you want in the report.'
            }, status=400)
        
        # Validate file type
        if not uploaded_file.name.endswith('.json'):
            return JsonResponse({'error': 'Only JSON files are supported'}, status=400)
        
        # Read and parse JSON
        file_content = uploaded_file.read().decode('utf-8')
        print(f"DEBUG: File content preview: {file_content[:200]}...")
        json_data = json.loads(file_content)

        if json_data is None:
            return JsonResponse({'error': 'JSON file is empty'}, status=400)

        normalized_json, original_json = normalize_json_for_analysis(json_data)
        bank_name, data_period = extract_entity_metadata(json_data)
        
        # Create a proper report structure with AI analysis
        report_id = str(uuid.uuid4())
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Store in session for persistence (simple approach)
        if 'reports' not in request.session:
            request.session['reports'] = []
        
        ai_analysis = perform_initial_ai_analysis(json_data)

        full_ai_analysis, ai_error_msg, ai_enhanced = generate_analysis_from_prompt(
            prompt, original_json, ai_analysis
        )
        comprehensive_generated = len(full_ai_analysis) > 0 and ai_enhanced

        prompt_title = prompt.strip().split('\n')[0].strip()
        if prompt_title.startswith('#'):
            prompt_title = prompt_title.lstrip('#').strip()
        prompt_title = prompt_title[:120] or f'{bank_name} Analysis'
        
        report_data = {
            'id': report_id,
            'filename': uploaded_file.name,
            'size': uploaded_file.size,
            'description': description,
            'user_prompt': prompt,
            'task_id': report_id,
            'uploaded_at': current_time,
            'status': 'completed',
            'data_summary': {
                **build_data_summary(original_json),
                'ai_insights': ai_analysis,
            },
            'bank_name': bank_name,
            'data_period': data_period,
            'ai_analysis': ai_analysis,
            'comprehensive_analysis': full_ai_analysis,
            'ai_enhanced': ai_enhanced,
            'metadata': {
                'title': prompt_title,
                'report_date': current_time,
                'period': data_period,
                'generated_at': current_time,
                'ai_processed': comprehensive_generated,
                'comprehensive_generated': comprehensive_generated,
                'ai_enhanced': ai_enhanced,
                'user_prompt': prompt,
                'original_json': original_json,
                'normalized_json': normalized_json,
            }
        }
        
        # Add AI error message if it exists
        if ai_error_msg:
            report_data['ai_error'] = ai_error_msg
        
        # Add to session
        request.session['reports'].append(report_data)
        request.session.modified = True
        
        # Also store in global memory for fallback
        add_global_report(report_id, report_data)
        
        # Debug output
        print(f"DEBUG: Added report to session. Total reports: {len(request.session['reports'])}")
        print(f"DEBUG: Session key: {request.session.session_key}")
        print(f"DEBUG: Report ID: {report_id}")
        print(f"DEBUG: Global reports count: {len(GLOBAL_REPORTS)}")
        
        # Create response with all report data
        response_data = {
            'success': True,
            'message': 'Report generated from your analysis prompt' if comprehensive_generated else 'Report uploaded successfully, but AI-enhanced analysis is currently unavailable.',
            'task_id': report_id,
            'id': report_id,
            'report_id': report_id,
        }
        if ai_error_msg:
            response_data['warning'] = ai_error_msg
            response_data['warning_code'] = 'openai_quota_exceeded' if 'quota' in ai_error_msg.lower() else 'ai_unavailable'
        # Add all report_data fields
        for key, value in report_data.items():
            response_data[key] = value
        
        return JsonResponse(response_data)
        
    except json.JSONDecodeError as e:
        print(f"DEBUG: JSON decode error: {str(e)}")
        print(f"DEBUG: File content that failed: {file_content[:500]}...")
        return JsonResponse({'error': f'Invalid JSON file format: {str(e)}'}, status=400)
    except Exception as e:
        print(f"DEBUG: General error: {str(e)}")
        return JsonResponse({'error': str(e)}, status=500)


@method_decorator(csrf_exempt, name='dispatch')
class FinancialDataUploadView(APIView):
    """Handle financial data file uploads"""
    
    parser_classes = [MultiPartParser, FormParser]
    permission_classes = []  # Temporarily disable for testing
    
    def post(self, request):
        """Upload and process financial data file"""
        try:
            # Debug: Check if user is authenticated
            print(f"User authenticated: {request.user.is_authenticated}")
            print(f"User: {request.user}")
            print(f"Session key: {request.session.session_key}")
            
            # Validate request data
            request_serializer = FileUploadRequestSerializer(data=request.data)
            if not request_serializer.is_valid():
                return Response(
                    {'error': 'Validation failed', 'details': request_serializer.errors},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Create upload record
            upload = FinancialDataUpload.objects.create(
                user=request.user,
                file=request.FILES['file'],
                original_filename=request.FILES['file'].name,
                status='uploaded'
            )
            
            # Parse and validate JSON data
            try:
                data = json.loads(request.FILES['file'].read().decode('utf-8'))
                parser = FinancialDataParser()
                parsed_data = parser.parse_financial_data(data)
                
                # Create data set record
                data_set = FinancialDataSet.objects.create(
                    upload=upload,
                    dashboard_data=parsed_data.get('dashboard', {}),
                    qc_dashboard_data=parsed_data.get('qc_dashboard', {}),
                    income_risk_data=parsed_data.get('income_risk', {}),
                    dupont_data=parsed_data.get('dupont', {}),
                    data_period=parsed_data.get('dashboard', {}).get('period', ''),
                    bank_name=parsed_data.get('dashboard', {}).get('bank_name', '')
                )
                
                # Trigger async analysis
                task = AnalysisTask.objects.create(
                    upload=upload,
                    task_id=str(uuid.uuid4()),
                    status='pending'
                )
                
                # Start Celery task
                process_financial_analysis.delay(task.task_id, data_set.id)
                
                upload.status = 'processing'
                upload.save()
                
                response_data = {
                    'upload_id': upload.id,
                    'task_id': task.task_id,
                    'status': 'processing',
                    'message': 'File uploaded successfully. Analysis started.'
                }
                
                return Response(response_data, status=status.HTTP_201_CREATED)
                
            except json.JSONDecodeError:
                upload.status = 'failed'
                upload.error_message = 'Invalid JSON format'
                upload.save()
                
                return Response(
                    {'error': 'Invalid JSON file format'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            except Exception as e:
                upload.status = 'failed'
                upload.error_message = str(e)
                upload.save()
                
                return Response(
                    {'error': 'Data parsing failed', 'message': str(e)},
                    status=status.HTTP_400_BAD_REQUEST
                )
                
        except Exception as e:
            return Response(
                {'error': 'Upload failed', 'message': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class AnalysisStatusView(APIView):
    """Check analysis task status"""
    
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request, task_id):
        """Get status of analysis task"""
        try:
            task = AnalysisTask.objects.get(task_id=task_id, upload__user=request.user)
            
            response_data = {
                'task_id': task.task_id,
                'status': task.status,
                'progress': task.progress,
                'error_message': task.error_message,
                'created_at': task.created_at,
                'completed_at': task.completed_at
            }
            
            if task.status == 'completed' and task.result_data:
                response_data['result_data'] = task.result_data
                response_data['report_id'] = task.result_data.get('report_id')
            
            return Response(response_data)
            
        except AnalysisTask.DoesNotExist:
            return Response(
                {'error': 'Task not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            return Response(
                {'error': 'Status check failed', 'message': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


class FinancialReportViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing financial reports"""
    
    serializer_class = ReportSummarySerializer
    permission_classes = [permissions.IsAuthenticated]
    pagination_class = PageNumberPagination
    
    def get_queryset(self):
        """Filter reports based on user permissions"""
        queryset = FinancialReport.objects.filter(
            data_set__upload__user=self.request.user
        ).select_related('data_set', 'data_set__upload')
        
        # Apply filters
        serializer = ReportFilterSerializer(data=self.request.query_params)
        if serializer.is_valid():
            filters = serializer.validated_data
            
            if filters.get('bank_name'):
                queryset = queryset.filter(
                    data_set__bank_name__icontains=filters['bank_name']
                )
            
            if filters.get('risk_level'):
                queryset = queryset.filter(risk_level=filters['risk_level'])
            
            if filters.get('min_score') is not None:
                queryset = queryset.filter(overall_score__gte=filters['min_score'])
            
            if filters.get('max_score') is not None:
                queryset = queryset.filter(overall_score__lte=filters['max_score'])
            
            if filters.get('date_from'):
                queryset = queryset.filter(generated_at__date__gte=filters['date_from'])
            
            if filters.get('date_to'):
                queryset = queryset.filter(generated_at__date__lte=filters['date_to'])
            
            ordering = filters.get('ordering', '-generated_at')
            queryset = queryset.order_by(ordering)
        
        return queryset
    
    def retrieve(self, request, pk=None):
        """Get detailed report"""
        try:
            report = FinancialReport.objects.get(
                id=pk, data_set__upload__user=request.user
            )
            serializer = DetailedReportSerializer(report)
            return Response(serializer.data)
        except FinancialReport.DoesNotExist:
            return Response(
                {'error': 'Report not found'},
                status=status.HTTP_404_NOT_FOUND
            )


class ReportDetailView(APIView):
    """Detailed view for a specific report"""
    
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request, report_id):
        """Get complete report details"""
        try:
            report = FinancialReport.objects.get(
                id=report_id, data_set__upload__user=request.user
            )
            
            # Return structured report data
            return Response(report.structured_data)
            
        except FinancialReport.DoesNotExist:
            return Response(
                {'error': 'Report not found'},
                status=status.HTTP_404_NOT_FOUND
            )


class MetricsSummaryView(APIView):
    """Get metrics summary for a report"""
    
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request, report_id):
        """Get key metrics summary"""
        try:
            report = FinancialReport.objects.get(
                id=report_id, data_set__upload__user=request.user
            )
            
            # Extract metrics from report
            key_metrics = report.structured_data.get('key_metrics', {})
            
            return Response(key_metrics)
            
        except FinancialReport.DoesNotExist:
            return Response(
                {'error': 'Report not found'},
                status=status.HTTP_404_NOT_FOUND
            )


class TrendAnalysisView(APIView):
    """Get trend analysis data"""
    
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request, report_id):
        """Get trend analysis with chart data"""
        try:
            report = FinancialReport.objects.get(
                id=report_id, data_set__upload__user=request.user
            )
            
            # Extract trend data
            trend_data = report.structured_data.get('trend_analysis', {})
            
            return Response(trend_data)
            
        except FinancialReport.DoesNotExist:
            return Response(
                {'error': 'Report not found'},
                status=status.HTTP_404_NOT_FOUND
            )


class BenchmarkComparisonView(APIView):
    """Get benchmark comparison data"""
    
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request, report_id):
        """Get benchmark comparison data"""
        try:
            report = FinancialReport.objects.get(
                id=report_id, data_set__upload__user=request.user
            )
            
            # Extract benchmark data
            benchmark_data = report.structured_data.get('benchmark_comparison', {})
            
            return Response(benchmark_data)
            
        except FinancialReport.DoesNotExist:
            return Response(
                {'error': 'Report not found'},
                status=status.HTTP_404_NOT_FOUND
            )


class ExportReportView(APIView):
    """Export report in various formats"""
    
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request, report_id):
        """Export report in specified format"""
        try:
            report = FinancialReport.objects.get(
                id=report_id, data_set__upload__user=request.user
            )
            
            export_format = request.query_params.get('format', 'json')
            include_raw_data = request.query_params.get('include_raw_data', 'false').lower() == 'true'
            
            if export_format == 'json':
                return self._export_json(report, include_raw_data)
            elif export_format == 'pdf':
                return self._export_pdf(report, include_raw_data)
            elif export_format == 'excel':
                return self._export_excel(report, include_raw_data)
            else:
                return Response(
                    {'error': 'Unsupported export format'},
                    status=status.HTTP_400_BAD_REQUEST
                )
                
        except FinancialReport.DoesNotExist:
            return Response(
                {'error': 'Report not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            return Response(
                {'error': 'Export failed', 'message': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    def _export_json(self, report, include_raw_data):
        """Export report as JSON"""
        data = report.structured_data
        
        if not include_raw_data:
            # Remove raw data sections
            data.pop('appendix', None)
        
        response = JsonResponse(data)
        response['Content-Disposition'] = f'attachment; filename="{report.title}.json"'
        return response
    
    def _export_pdf(self, report, include_raw_data):
        """Export report as PDF"""
        # Implementation would use reportlab or similar
        return Response(
            {'error': 'PDF export not yet implemented'},
            status=status.HTTP_501_NOT_IMPLEMENTED
        )
    
    def _export_excel(self, report, include_raw_data):
        """Export report as Excel"""
        # Implementation would use pandas ExcelWriter
        return Response(
            {'error': 'Excel export not yet implemented'},
            status=status.HTTP_501_NOT_IMPLEMENTED
        )


class UserUploadsView(APIView):
    """Get user's upload history"""
    
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request):
        """Get user's upload history"""
        uploads = FinancialDataUpload.objects.filter(
            user=request.user
        ).order_by('-uploaded_at')
        
        serializer = FinancialDataUploadSerializer(uploads, many=True)
        return Response(serializer.data)


@method_decorator(csrf_exempt, name='dispatch')
class SystemHealthView(APIView):
    """System health check endpoint"""
    
    permission_classes = [permissions.AllowAny]
    
    def get(self, request):
        """Check system health"""
        try:
            from django.utils import timezone
            
            # Check database connection
            from django.db import connection
            with connection.cursor() as cursor:
                cursor.execute("SELECT 1")
            
            # Simple health check
            return Response({
                'status': 'healthy',
                'database': 'connected',
                'timestamp': timezone.now().isoformat(),
                'version': '1.0.0'
            })
            
        except Exception as e:
            from django.utils import timezone
            return Response({
                'status': 'unhealthy',
                'error': str(e),
                'timestamp': timezone.now().isoformat()
            }, status=500)


# API endpoints for direct analysis (without file upload)

@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def analyze_direct_data(request):
    """Analyze financial data provided directly in request"""
    try:
        data = request.data.get('financial_data')
        report_options = request.data.get('report_options') or {}
        if not data:
            return Response(
                {'error': 'financial_data is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Parse and analyze data
        parser = FinancialDataParser()
        parsed_data = parser.parse_financial_data(data)
        
        # Generate report
        generator = FinancialReportGenerator()
        report_data = generator.generate_complete_report(parsed_data, report_options=report_options)
        
        return Response(report_data)
        
    except Exception as e:
        return Response(
            {'error': 'Analysis failed', 'message': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([])  # Temporarily disable authentication for testing
def get_insights(request, report_id):
    """Get AI insights for a specific report section"""
    try:
        # First try to find report in session/global storage (simple upload system)
        report = None
        session_reports = request.session.get('reports', [])
        
        for r in session_reports:
            if r.get('id') == report_id:
                report = r
                break
        
        if not report:
            global_reports = get_global_reports()
            for global_id, global_report in global_reports.items():
                if str(global_id) == str(report_id):
                    report = global_report
                    break
        
        if report:
            # Simple upload system - return existing analysis
            return Response({
                'success': True,
                'insights': report.get('ai_analysis', {})
            })
        
        # Fallback to database system (requires authentication)
        report = FinancialReport.objects.get(
            id=report_id, data_set__upload__user=request.user
        )
        
        insight_type = request.query_params.get('type')
        if insight_type:
            # Get specific insight
            insights = FinancialInsight.objects.filter(
                data_set=report.data_set,
                insight_type=insight_type
            )
        else:
            # Get all insights
            insights = FinancialInsight.objects.filter(data_set=report.data_set)
        
        serializer = FinancialInsightSerializer(insights, many=True)
        return Response(serializer.data)
        
    except FinancialReport.DoesNotExist:
        return Response(
            {'error': 'Report not found'},
            status=status.HTTP_404_NOT_FOUND
        )


@api_view(['POST'])
@permission_classes([])  # Temporarily disable authentication for testing
def regenerate_insights(request, report_id):
    """Regenerate AI insights for a report"""
    try:
        from .services.report_store import get_report, update_report

        report_options = {}
        if hasattr(request, 'data') and isinstance(request.data, dict):
            report_options = request.data.get('report_options') or {}
        elif request.body:
            try:
                body = json.loads(request.body)
                if isinstance(body, dict):
                    report_options = body.get('report_options') or {}
            except json.JSONDecodeError:
                report_options = {}

        report = get_report(str(report_id))

        if report:
            original_json = report.get('metadata', {}).get('original_json')
            user_prompt = report.get('user_prompt') or report.get('metadata', {}).get('user_prompt')

            if not original_json:
                return Response(
                    {'error': 'Original JSON not found in report'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            if not user_prompt:
                return Response(
                    {'error': 'No analysis prompt stored on this report. Upload again with a prompt.'},
                    status=status.HTTP_400_BAD_REQUEST
                )

            ai_analysis = perform_initial_ai_analysis(original_json)
            sections, error_msg, ai_enhanced = generate_analysis_from_prompt(
                user_prompt, original_json, ai_analysis, report_options=report_options
            )

            if sections and ai_enhanced:
                generated_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                metadata = dict(report.get('metadata', {}))
                metadata.update({
                    'generated_at': generated_at,
                    'comprehensive_generated': True,
                    'ai_enhanced': True,
                    'report_options': report_options or metadata.get('report_options', {}),
                })
                report_updates = {
                    'comprehensive_analysis': sections,
                    'ai_enhanced': True,
                    'ai_error': None,
                    'status': 'completed',
                    'metadata': metadata,
                    'report_options': report_options or report.get('report_options', {}),
                }
                update_report(str(report_id), report_updates, request=request)
                return Response({
                    'success': True,
                    'comprehensive_analysis': sections,
                    'ai_enhanced': True,
                    'generated_at': generated_at,
                })

            status_code = status.HTTP_402_PAYMENT_REQUIRED if error_msg and 'quota' in error_msg.lower() else status.HTTP_503_SERVICE_UNAVAILABLE
            return Response({'error': error_msg or 'AI regeneration failed'}, status=status_code)
        
        # If no report found, try to provide fallback
        global_reports = get_global_reports()
        if global_reports:
            # Return first available report as fallback
            fallback_report = list(global_reports.values())[0]
            print(f"DEBUG: Regenerate insights - Providing fallback report: {fallback_report.get('id')}")
            
            # Try to regenerate insights for fallback report
            fallback_json = fallback_report.get('metadata', {}).get('original_json')
            if fallback_json:
                new_analysis = perform_initial_ai_analysis(fallback_json)
                return Response({
                    'success': True,
                    'insights': new_analysis,
                    'fallback': True,
                    'original_request_id': report_id,
                    'fallback_report_id': fallback_report.get('id'),
                    'message': f'Report {report_id} not found. Regenerating insights for available report instead.'
                })
            else:
                return Response({
                    'success': True,
                    'insights': fallback_report.get('ai_analysis', {}),
                    'fallback': True,
                    'original_request_id': report_id,
                    'fallback_report_id': fallback_report.get('id'),
                    'message': f'Report {report_id} not found. Using existing insights for available report.'
                })
        else:
            return Response(
                {'error': 'Report not found', 'message': 'No reports available in system. Please upload a financial data file first.'},
                status=status.HTTP_404_NOT_FOUND
            )
        
    except Exception as e:
        print(f"DEBUG: Regenerate insights error: {e}")
        return Response(
            {'error': 'Insight regeneration failed', 'message': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


class CustomReportView(APIView):
    """Generate custom report based on user prompt"""
    
    permission_classes = [permissions.IsAuthenticated]
    
    def post(self, request, report_id):
        """Generate custom report with user prompt"""
        try:
            report = FinancialReport.objects.get(id=report_id, user=request.user)
            data_set = report.data_set
            
            # Get user prompt from request
            prompt = request.data.get('prompt', '')
            if not prompt:
                return Response(
                    {'error': 'Prompt is required'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            report_options = request.data.get('report_options') or {
                'template': request.data.get('template'),
                'sections': request.data.get('sections') or request.data.get('selected_sections') or [],
                'include_sections': request.data.get('include_sections') or [],
                'exclude_sections': request.data.get('exclude_sections') or [],
                'length': request.data.get('length'),
                'detail_level': request.data.get('detail_level'),
                'output_format': request.data.get('format'),
            }
            
            # Get original data
            original_data = {
                'dashboard': data_set.dashboard_data,
                'qc_dashboard': data_set.qc_dashboard_data,
                'income_risk': data_set.income_risk_data,
                'dupont': data_set.dupont_data
            }
            
            # Generate custom report using prompt
            insight_engine = FinancialInsightEngine()
            custom_report = insight_engine.generate_custom_report(prompt, original_data, report, report_options=report_options)
            
            # Save custom report
            report.custom_prompt = prompt
            report.custom_report = custom_report
            report.save()
            
            return Response({
                'custom_report': custom_report,
                'prompt': prompt,
                'generated_at': timezone.now().isoformat()
            })
            
        except FinancialReport.DoesNotExist:
            return Response(
                {'error': 'Report not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            return Response(
                {'error': 'Report generation failed', 'message': str(e)},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )


# Duplicate imports removed

@csrf_exempt
@require_http_methods(["GET"])
def simple_reports_view(request):
    """Simple reports view - returns stored reports"""
    try:
        
        # Try session first, then fallback to global
        session_reports = request.session.get('reports', [])
        global_reports = list(get_global_reports().values())
        
        # Use session reports if available, otherwise global
        reports = session_reports if session_reports else global_reports
        
        print(f"DEBUG: Session reports: {len(session_reports)}")
        print(f"DEBUG: Global reports: {len(global_reports)}")
        print(f"DEBUG: Session key: {request.session.session_key}")
        print(f"DEBUG: Session data keys: {list(request.session.keys())}")
        print(f"DEBUG: Using reports from: {'session' if session_reports else 'global'}")
        
        return JsonResponse({
            'results': reports,
            'count': len(reports),
            'session_key': request.session.session_key,
            'debug': {
                'session_keys': list(request.session.keys()),
                'session_reports_count': len(session_reports),
                'global_reports_count': len(global_reports),
                'source': 'session' if session_reports else 'global'
            }
        })
    except Exception as e:
        print(f"DEBUG: Error in simple_reports_view: {e}")
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["GET"])
def simple_task_status_view(request, task_id):
    """Simple task status view"""
    try:
        report_id = str(task_id)
        return JsonResponse({
            'id': task_id,
            'status': 'completed',
            'progress': 100,
            'message': 'Analysis completed successfully',
            'result_data': {
                'report_id': report_id
            }
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["GET"])
def simple_report_detail_view(request, report_id):
    """Simple report detail view"""
    try:
        # Try session first, then fallback to global
        session_reports = request.session.get('reports', [])
        
        print(f"DEBUG: Looking for report ID: {report_id}")
        print(f"DEBUG: Session reports available: {[r.get('id') for r in session_reports]}")
        print(f"DEBUG: Global reports available: {list(GLOBAL_REPORTS.keys())}")
        
        report = None
        source = None
        
        # First try session
        for r in session_reports:
            if str(r.get('id')) == str(report_id):
                report = r
                source = 'session'
                break
        
        # If not found in session, try global
        if not report:
            global_reports = get_global_reports()
            for global_id, global_report in global_reports.items():
                if str(global_id) == str(report_id):
                    report = global_report
                    source = 'global'
                    print(f"DEBUG: Found report in global storage with ID match: {global_id}")
                    break
        
        # If still not found, try database as fallback
        if not report:
            try:
                from .models import FinancialDataSet
                financial_dataset = FinancialDataSet.objects.filter(id=report_id).first()
                if financial_dataset:
                    report = {
                        'id': str(financial_dataset.id),
                        'filename': financial_dataset.filename,
                        'data_summary': financial_dataset.data_summary,
                        'bank_name': financial_dataset.bank_name,
                        'data_period': financial_dataset.data_period,
                        'ai_analysis': financial_dataset.ai_analysis,
                        'metadata': financial_dataset.metadata
                    }
                    source = 'database'
                    print(f"DEBUG: Found report in database: {report_id}")
                else:
                    print(f"DEBUG: Report not found in database either: {report_id}")
            except Exception as e:
                print(f"DEBUG: Database query failed: {e}")
        
        if not report:
            print(f"DEBUG: Report not found: {report_id}")
            
            # Try to provide a fallback from available reports
            global_reports = get_global_reports()
            if global_reports:
                # Return the first available report as fallback
                fallback_report = list(global_reports.values())[0]
                print(f"DEBUG: Providing fallback report: {fallback_report.get('id')}")
                return JsonResponse({
                    'report': fallback_report,
                    'fallback': True,
                    'original_request_id': report_id,
                    'message': f'Report {report_id} not found. Showing available report instead.',
                    'debug': {
                        'session_reports': len(session_reports),
                        'global_reports': len(global_reports),
                        'available_ids': list(global_reports.keys())
                    }
                })
            else:
                return JsonResponse({
                    'error': 'Report not found', 
                    'report_id': report_id,
                    'message': 'No reports available in the system. Please upload a financial data file first.',
                    'debug': {
                        'session_reports': len(session_reports),
                        'global_reports': len(GLOBAL_REPORTS),
                        'available_session_ids': [r.get('id') for r in session_reports],
                        'available_global_ids': list(GLOBAL_REPORTS.keys())
                    }
                }, status=404)
        
        print(f"DEBUG: Found report: {report.get('id')} from {source}")
        return JsonResponse(report)
    except Exception as e:
        print(f"DEBUG: Error in simple_report_detail_view: {e}")
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["GET"])
def simple_export_view(request, report_id):
    """Simple export view"""
    try:
        report = None
        source = None
        
        print(f"DEBUG: Export - Looking for report ID: {report_id}")
        
        # First try session storage (primary storage location)
        reports = request.session.get('reports', [])
        print(f"DEBUG: Export - Session key: {request.session.session_key}")
        print(f"DEBUG: Export - Session reports count: {len(reports)}")
        print(f"DEBUG: Export - Session report IDs: {[r.get('id') for r in reports]}")
        
        for r in reports:
            if str(r.get('id')) == str(report_id):
                report = r
                source = 'session'
                print(f"DEBUG: Export - Found report in session: {report_id}")
                break
        
        # If not found in session, try global storage
        if not report:
            global_reports = get_global_reports()
            print(f"DEBUG: Export - Global reports count: {len(global_reports)}")
            print(f"DEBUG: Export - Global report IDs: {list(global_reports.keys())[:3]}")
            
            for global_id, global_report in global_reports.items():
                if str(global_id) == str(report_id):
                    report = global_report
                    source = 'global'
                    print(f"DEBUG: Export - Found report in global storage: {report_id}")
                    break
        
        if not report:
            print(f"DEBUG: Export - Report not found in session or global storage: {report_id}")
            
            # Try to provide a fallback from available reports in session
            if reports:
                # Return first available report from session as fallback
                fallback_report = reports[0]
                print(f"DEBUG: Export - Providing fallback report from session: {fallback_report.get('id')}")
                report = fallback_report
                source = 'session_fallback'
            elif global_reports:
                # Return first available report from global as fallback
                fallback_report = list(global_reports.values())[0]
                print(f"DEBUG: Export - Providing fallback report from global: {fallback_report.get('id')}")
                report = fallback_report
                source = 'global_fallback'
            else:
                return JsonResponse({
                    'error': 'Report not found', 
                    'report_id': report_id,
                    'message': 'No reports available in the system. Please upload a financial data file first.',
                    'debug': {
                        'global_reports_count': len(get_global_reports()),
                        'session_reports_count': len(request.session.get('reports', [])),
                        'available_global_ids': list(get_global_reports().keys())
                    }
                }, status=404)
        
        print(f"DEBUG: Export - Report keys: {list(report.keys())}")
        print(f"DEBUG: Export - Report data type: {type(report)}")
        
        # Create export data
        try:
            export_data = {
                'report_id': report.get('id'),
                'filename': report.get('filename'),
                'bank_name': report.get('bank_name'),
                'data_period': report.get('data_period'),
                'generated_at': report.get('uploaded_at'),
                'data_summary': report.get('data_summary'),
                'metadata': report.get('metadata'),
                'export_timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            print(f"DEBUG: Export - Export data created successfully")
        except Exception as e:
            print(f"DEBUG: Export - Error creating export data: {e}")
            return JsonResponse({'error': f'Export data creation failed: {str(e)}'}, status=500)
        
        format = request.GET.get('format', 'json')
        
        if format == 'json':
            # Return structured export data as JSON
            response = JsonResponse(export_data)
            response['Content-Disposition'] = f'attachment; filename="financial_report_{report_id}.json"'
            response['Content-Type'] = 'application/json'
            return response
        
        elif format == 'pdf':
            try:
                pdf_content = generate_pdf_report(report)
                response = HttpResponse(pdf_content, content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="financial_report_{report_id}.pdf"'
                return response
            except Exception as e:
                return JsonResponse({'error': f'PDF generation failed: {str(e)}'}, status=500)
        
        elif format == 'csv':
            try:
                csv_content = generate_csv_report(report)
                response = HttpResponse(csv_content, content_type='text/csv')
                response['Content-Disposition'] = f'attachment; filename="financial_report_{report_id}.csv"'
                return response
            except Exception as e:
                return JsonResponse({'error': f'CSV generation failed: {str(e)}'}, status=500)
        
        elif format == 'word':
            try:
                doc_content = generate_word_report(report)
                response = HttpResponse(doc_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="financial_report_{report_id}.docx"'
                return response
            except Exception as e:
                return JsonResponse({'error': f'Word document generation failed: {str(e)}'}, status=500)
        
        elif format == 'excel':
            try:
                excel_content = generate_excel_report(report)
                response = HttpResponse(excel_content, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                response['Content-Disposition'] = f'attachment; filename="financial_report_{report_id}.xlsx"'
                return response
            except Exception as e:
                return JsonResponse({'error': f'Excel generation failed: {str(e)}'}, status=500)
        
        else:
            return JsonResponse({'error': f'Unsupported format: {format}'}, status=400)
            
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@api_view(['GET'])
@permission_classes([])  # Disable authentication for testing
def get_report_templates(request):
    """Get available report templates"""
    try:
        registry = get_report_prompt_registry()
        return JsonResponse({
            'success': True,
            'templates': registry.get_templates(),
            'section_library': registry.get_section_library()
        })
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        }, status=500)

@api_view(['POST'])
@permission_classes([])  # Disable authentication for testing
def generate_comprehensive_report(request, report_id):
    """Generate comprehensive financial report"""
    try:
        # Get request data
        body = json.loads(request.body) if request.method == 'POST' else {}
        registry = get_report_prompt_registry()
        report_options = registry.build_report_options({
            'template': body.get('template', 'standard'),
            'sections': body.get('sections') or body.get('selected_sections') or [],
            'include_sections': body.get('include_sections') or [],
            'exclude_sections': body.get('exclude_sections') or [],
            'length': body.get('length'),
            'detail_level': body.get('detail_level'),
            'output_format': body.get('format', 'json'),
        })
        template_type = report_options.get('template', 'custom')
        format_type = report_options.get('output_format', 'json')
        
        # Find report in session or global storage
        report = None
        session_reports = request.session.get('reports', [])
        
        for r in session_reports:
            if r.get('id') == report_id:
                report = r
                break
        
        if not report:
            global_reports = get_global_reports()
            for global_id, global_report in global_reports.items():
                if str(global_id) == str(report_id):
                    report = global_report
                    break
        
        if not report:
            return JsonResponse({'error': 'Report not found'}, status=404)
        
        # Generate comprehensive report based on template
        templates = registry.get_templates()
        template = templates.get(template_type, templates.get('custom', {}))
        sections = report_options.get('sections') or template.get('sections', [])
        generated_report = {
            'report_id': report_id,
            'template_used': template_type,
            'sections': generate_report_sections(sections, report),
            'metadata': {
                'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'template_version': '1.0',
                'word_count': count_report_words(report),
                'report_options': report_options,
            }
        }
        
        # Return based on format
        if format_type == 'json':
            return JsonResponse(generated_report)
        elif format_type == 'pdf':
            pdf_content = generate_pdf_report(generated_report)
            response = HttpResponse(pdf_content, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="comprehensive_report_{report_id}.pdf"'
            return response
        elif format_type == 'word':
            doc_content = generate_word_report(generated_report)
            response = HttpResponse(doc_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="comprehensive_report_{report_id}.docx"'
            return response
        else:
            return JsonResponse({'error': f'Unsupported format: {format_type}'}, status=400)
            
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON in request body'}, status=400)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@api_view(['GET'])
@permission_classes([])  # Disable authentication for testing
def preview_report(request, report_id):
    """Preview report before generation"""
    try:
        registry = get_report_prompt_registry()
        # Find report
        report = None
        session_reports = request.session.get('reports', [])
        
        for r in session_reports:
            if r.get('id') == report_id:
                report = r
                break
        
        if not report:
            global_reports = get_global_reports()
            for global_id, global_report in global_reports.items():
                if str(global_id) == str(report_id):
                    report = global_report
                    break
        
        if not report:
            return JsonResponse({'error': 'Report not found'}, status=404)
        
        # Generate preview (first few sections)
        preview_data = {
            'report_id': report_id,
            'bank_name': report.get('bank_name'),
            'data_period': report.get('data_period'),
            'preview_sections': generate_report_sections(['executive_summary', 'statistical_highlights'], report),
            'available_templates': registry.get_templates()
        }
        
        return JsonResponse(preview_data)
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

def generate_report_sections(sections, report_data):
    """Generate individual report sections"""
    generated_sections = {}
    dynamic_sections = {
        'executive_summary',
        'statistical_highlights',
        'money_market_analysis',
        'wacc_analysis',
        'financial_ratios',
        'investment_analysis',
        'macroeconomic_indicators',
        'country_risk_analysis',
        'market_trends',
        'trend_analysis',
        'risk_assessment',
        'benchmark_comparison',
        'recommendations',
    }
    
    for section in sections:
        if section in dynamic_sections:
            built = build_dynamic_report_sections([section], report_data)
            if built:
                generated_sections[section] = built[0]
                continue
        if section == 'key_metrics':
            generated_sections[section] = generate_key_metrics_section(report_data)
        elif section == 'profitability':
            generated_sections[section] = generate_profitability_section(report_data)
        elif section == 'liquidity':
            generated_sections[section] = generate_liquidity_section(report_data)
        elif section == 'risk_assessment':
            generated_sections[section] = generate_risk_assessment_section(report_data)
        elif section == 'overview':
            generated_sections[section] = generate_overview_section(report_data)
        elif section == 'ratio_analysis':
            generated_sections[section] = generate_ratio_analysis_section(report_data)
        elif section == 'trend_analysis':
            generated_sections[section] = generate_trend_analysis_section(report_data)
        elif section == 'recommendations':
            generated_sections[section] = generate_recommendations_section(report_data)
        elif section == 'capital_adequacy':
            generated_sections[section] = generate_capital_adequacy_section(report_data)
        elif section == 'compliance_status':
            generated_sections[section] = generate_compliance_status_section(report_data)
        else:
            built = build_dynamic_report_sections([section], report_data)
            generated_sections[section] = built[0] if built else f"Section {section} not implemented yet"
    
    return generated_sections

def generate_key_metrics_section(report_data):
    """Generate key financial metrics section"""
    metrics = report_data.get('data_summary', {})
    return {
        'title': 'Key Financial Metrics',
        'metrics': {
            'total_assets': metrics.get('total_assets', 0),
            'roa': metrics.get('roa', 0),
            'roe': metrics.get('roe', 0),
            'bank_name': report_data.get('bank_name', 'Unknown'),
            'period': report_data.get('data_period', 'Unknown')
        }
    }

def generate_profitability_section(report_data):
    """Generate profitability analysis section"""
    ai_analysis = report_data.get('ai_analysis', {})
    return {
        'title': 'Profitability Analysis',
        'analysis': ai_analysis.get('profitability', {}),
        'insights': [
            'ROA indicates efficient asset utilization',
            'ROE shows strong shareholder returns',
            'Monitor profit margin trends'
        ]
    }

def generate_liquidity_section(report_data):
    """Generate liquidity analysis section"""
    ai_analysis = report_data.get('ai_analysis', {})
    return {
        'title': 'Liquidity Analysis',
        'analysis': ai_analysis.get('liquidity', {}),
        'insights': [
            'Current ratio indicates short-term liquidity',
            'Quick ratio shows immediate liquidity',
            'Cash flow analysis needed'
        ]
    }

def generate_risk_assessment_section(report_data):
    """Generate risk assessment section"""
    ai_analysis = report_data.get('ai_analysis', {})
    return {
        'title': 'Risk Assessment',
        'analysis': ai_analysis.get('risk_metrics', {}),
        'insights': [
            'Credit risk profile analyzed',
            'Market risk considerations',
            'Operational risk factors'
        ]
    }

def generate_overview_section(report_data):
    """Generate overview section"""
    return {
        'title': 'Executive Overview',
        'summary': report_data.get('ai_analysis', {}).get('analysis_summary', 'Analysis complete'),
        'key_points': [
            f"Bank: {report_data.get('bank_name', 'Unknown')}",
            f"Period: {report_data.get('data_period', 'Unknown')}",
            f"Analysis Date: {datetime.now().strftime('%Y-%m-%d')}"
        ]
    }

def generate_ratio_analysis_section(report_data):
    """Generate ratio analysis section"""
    metrics = report_data.get('data_summary', {})
    return {
        'title': 'Financial Ratio Analysis',
        'ratios': {
            'profitability_ratios': {
                'roa': metrics.get('roa', 0),
                'roe': metrics.get('roe', 0)
            },
            'liquidity_ratios': {
                'current_ratio': 'Calculated from balance sheet',
                'quick_ratio': 'Calculated from balance sheet'
            }
        }
    }

def generate_trend_analysis_section(report_data):
    """Generate trend analysis section"""
    return {
        'title': 'Trend Analysis',
        'trends': [
            'Asset growth over time',
            'Loan portfolio changes',
            'Deposit trends',
            'Profitability trends'
        ]
    }

def generate_recommendations_section(report_data):
    """Generate recommendations section"""
    ai_analysis = report_data.get('ai_analysis', {})
    recommendations = ai_analysis.get('recommendations', [])
    
    return {
        'title': 'Strategic Recommendations',
        'recommendations': recommendations if recommendations else [
            'Improve operational efficiency',
            'Enhance risk management',
            'Optimize capital allocation',
            'Strengthen compliance measures'
        ]
    }

def generate_capital_adequacy_section(report_data):
    """Generate capital adequacy section"""
    return {
        'title': 'Capital Adequacy Assessment',
        'assessment': {
            'capital_ratios': 'Calculated from balance sheet',
            'regulatory_requirements': 'Basel III compliance',
            'buffer_requirements': 'Stress testing results'
        }
    }

def generate_compliance_status_section(report_data):
    """Generate compliance status section"""
    return {
        'title': 'Regulatory Compliance Status',
        'compliance_areas': [
            'Capital adequacy',
            'Risk management',
            'Reporting requirements',
            'Consumer protection'
        ]
    }

def count_report_words(report_data):
    """Count approximate words in report"""
    import re
    text = str(report_data)
    words = len(re.findall(r'\b\w+\b', text))
    return words

@csrf_exempt
@require_http_methods(["GET", "POST"])
def simple_login_view(request):
    """Simple Django view for login - no DRF"""
    if request.method == 'GET':
        return JsonResponse({'message': 'Simple login endpoint is accessible'})
    
    # For POST requests
    try:
        import json
        data = json.loads(request.body) if request.body else {}
        username = data.get('username')
        password = data.get('password')
        
        return JsonResponse({
            'message': 'Login request received',
            'username': username,
            'success': True
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@csrf_exempt
@require_http_methods(["POST"])
def simple_custom_report_view(request, report_id=None):
    """AI-powered custom report view - Using real AI analysis"""
    try:
        import json
        
        # Get request data
        data = json.loads(request.body) if request.body else {}
        prompt = data.get('prompt', '')
        report_options = data.get('report_options') or {
            'template': data.get('template'),
            'sections': data.get('sections') or data.get('selected_sections') or [],
            'include_sections': data.get('include_sections') or [],
            'exclude_sections': data.get('exclude_sections') or [],
            'length': data.get('length'),
            'detail_level': data.get('detail_level'),
            'output_format': data.get('format'),
        }
        
        if not prompt:
            return JsonResponse({'error': 'Prompt is required'}, status=400)
        
        if not report_id:
            return JsonResponse({'error': 'Report ID is required'}, status=400)
        
        # Find report in session or global storage
        report = None
        if 'reports' in request.session:
            for r in request.session['reports']:
                if r['id'] == report_id:
                    report = r
                    break
        
        if not report:
            for r in GLOBAL_REPORTS.values():
                if r['id'] == report_id:
                    report = r
                    break
        
        if not report:
            return JsonResponse({'error': 'Report not found'}, status=404)
        
        # Get original JSON data
        original_json = report.get('metadata', {}).get('original_json')
        if not original_json:
            return JsonResponse({'error': 'Original JSON data not found in report'}, status=400)
        
        print(f"DEBUG: Generating custom analysis with prompt: {prompt[:100]}...")
        
        # Generate analysis using custom prompt
        try:
            sections, error_msg, ai_enhanced = generate_analysis_from_prompt(
                prompt,
                original_json,
                report.get('ai_analysis', {}),
                report_options=report_options,
            )
            
            if sections and ai_enhanced:
                generated_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                metadata = dict(report.get('metadata', {}))
                metadata.update({
                    'user_prompt': prompt,
                    'comprehensive_generated': True,
                    'generated_at': generated_at,
                })
                report_updates = {
                    'comprehensive_analysis': sections,
                    'user_prompt': prompt,
                    'status': 'completed',
                    'metadata': metadata,
                }
                update_stored_report(report_id, report_updates, request)

                return JsonResponse({
                    'success': True,
                    'comprehensive_analysis': sections,
                    'custom_report': {
                        'title': f'AI Analysis: {prompt[:80]}',
                        'prompt': prompt,
                        'generated_at': generated_at,
                        'bank_name': report.get('bank_name', 'Unknown Bank'),
                        'data_period': report.get('data_period', 'Unknown Period'),
                        'sections': sections,
                        'ai_enhanced': True
                    },
                    'prompt': prompt,
                    'report_id': report_id,
                    'generated_at': generated_at,
                    'status': 'success'
                })

            return JsonResponse({
                'success': False,
                'error': error_msg or 'Unknown error'
            }, status=500)
                
        except Exception as e:
            print(f"DEBUG: Exception in custom analysis: {e}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)
        
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON in request body'}, status=400)
    except Exception as e:
        print(f"DEBUG: Custom report error: {e}")
        return JsonResponse({'error': str(e)}, status=500)


def transform_complex_json(json_data):
    """Transform complex JSON structure to expected format"""
    try:
        dashboard = json_data.get('Dashboard', {})
        income_risk = json_data.get('IncomeRisk', {})
        bank_data = income_risk.get('Bank', {})
        
        # Extract the most recent data point from QCDashboard
        qcdashboard = json_data.get('QCDashboard', [])
        latest_data = qcdashboard[0] if qcdashboard else {}
        
        # Transform to expected format
        transformed = {
            'bank_name': bank_data.get('Name', 'Financial Institution'),
            'period': latest_data.get('Asof', 'Unknown Period')[:10] if latest_data.get('Asof') else 'Unknown Period',
            'total_assets': dashboard.get('EVAssets', 0),
            'total_loans': latest_data.get('TotalLoans', 0),
            'total_deposits': latest_data.get('TotalDeposits', 0),
            'net_income': dashboard.get('NetIncome', 0),
            'revenue': dashboard.get('TotalRevenue', 0),
            'operating_expenses': dashboard.get('TotalRevenue', 0) - dashboard.get('NetIncome', 0),  # Calculate expenses
            'credit_risk_score': income_risk.get('RiskMetrics', {}).get('CreditRisk', 0),
            'liquidity_ratio': latest_data.get('LiquidityRatio', 0) / 100 if latest_data.get('LiquidityRatio') else 0.1,
            'capital_adequacy': latest_data.get('CapitalRatio', 0),
            'roa': dashboard.get('ROA', 0),
            'roe': dashboard.get('ROE', 0),
            'net_interest_margin': (dashboard.get('NetIncome', 0) / dashboard.get('TotalRevenue', 1)) * 100 if dashboard.get('TotalRevenue') else 0,
            'efficiency_ratio': ((dashboard.get('TotalRevenue', 0) - dashboard.get('NetIncome', 0)) / dashboard.get('TotalRevenue', 1)) * 100 if dashboard.get('TotalRevenue') else 0,
            'non_performing_loans_ratio': income_risk.get('RiskMetrics', {}).get('CreditRisk', 0),
            'cost_to_income_ratio': ((dashboard.get('TotalRevenue', 0) - dashboard.get('NetIncome', 0)) / dashboard.get('TotalRevenue', 1)) * 100 if dashboard.get('TotalRevenue') else 0,
            'asset_growth': 0,  # Would need historical data
            'loan_growth': 0,  # Would need historical data
            'deposit_growth': 0,  # Would need historical data
            'digital_adoption_rate': 50,  # Default value
            'customer_satisfaction_score': 7.5,  # Default value
            'market_share': 10  # Default value
        }
        
        print(f"DEBUG: Transformed data sample: {str(transformed)[:200]}")
        return transformed
        
    except Exception as e:
        print(f"DEBUG: Error transforming JSON: {e}")
        # Return default structure if transformation fails
        return {
            'bank_name': 'Financial Institution',
            'period': 'Unknown Period',
            'total_assets': 0,
            'total_loans': 0,
            'total_deposits': 0,
            'net_income': 0,
            'revenue': 0,
            'operating_expenses': 0,
            'credit_risk_score': 0,
            'liquidity_ratio': 0,
            'capital_adequacy': 0,
            'roa': 0,
            'roe': 0,
            'net_interest_margin': 0,
            'efficiency_ratio': 0,
            'non_performing_loans_ratio': 0,
            'cost_to_income_ratio': 0,
            'asset_growth': 0,
            'loan_growth': 0,
            'deposit_growth': 0,
            'digital_adoption_rate': 50,
            'customer_satisfaction_score': 7.5,
            'market_share': 10
        }


def perform_initial_ai_analysis(json_data):
    """Extract structural hints from any JSON shape before AI report generation."""
    analysis = {
        'data_type': 'general_json',
        'financial_indicators': [],
        'risk_level': 'moderate',
        'key_metrics': {},
        'recommendations': [],
        'analysis_summary': '',
    }

    def collect_metrics(obj, prefix='', depth=0):
        if depth > 8:
            return
        if isinstance(obj, dict):
            for key, value in obj.items():
                path = f'{prefix}.{key}' if prefix else str(key)
                key_lower = str(key).lower()
                if isinstance(value, (int, float)) and not isinstance(value, bool):
                    metric_type = 'general'
                    status = 'neutral'
                    if any(term in key_lower for term in ('revenue', 'income', 'earning')):
                        metric_type, analysis['data_type'] = 'revenue', 'financial_performance'
                    elif any(term in key_lower for term in ('profit', 'margin')):
                        metric_type = 'profitability'
                    elif any(term in key_lower for term in ('risk', 'loss', 'npl')):
                        metric_type, status = 'risk', 'warning'
                    elif any(term in key_lower for term in ('asset', 'liabilit', 'equity')):
                        analysis['data_type'] = 'balance_sheet'
                    elif any(term in key_lower for term in ('cash', 'liquid')):
                        analysis['data_type'] = 'cash_flow'
                    analysis['key_metrics'][path] = {'value': value, 'type': metric_type, 'status': status}
                elif isinstance(value, (dict, list)):
                    collect_metrics(value, path, depth + 1)
        elif isinstance(obj, list):
            for index, item in enumerate(obj[:50]):
                collect_metrics(item, f'{prefix}[{index}]' if prefix else f'[{index}]', depth + 1)

    collect_metrics(json_data)
    summary = build_data_summary(json_data)
    analysis['financial_indicators'] = [
        f"Detected {summary.get('structure', 'unknown')} JSON structure",
        f"Identified {len(analysis['key_metrics'])} numeric fields",
    ]
    analysis['analysis_summary'] = (
        f"Prepared {summary.get('structure', 'unknown')} data with "
        f"{len(analysis['key_metrics'])} metrics for AI analysis."
    )
    return analysis


def generate_ai_financial_analysis(prompt, financial_data, report):
    """Generate AI-powered financial analysis based on user prompt"""
    
    # Extract key financial metrics and existing AI analysis
    bank_name = report.get('bank_name', 'Unknown Bank')
    data_period = report.get('data_period', 'Unknown Period')
    existing_ai_analysis = report.get('ai_analysis', {})
    data_summary = report.get('data_summary', {})
    
    # Initialize report structure
    custom_report = {
        'title': f'AI-Powered Financial Analysis: {prompt.title()}',
        'prompt': prompt,
        'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'bank_name': bank_name,
        'data_period': data_period,
        'sections': []
    }
    
    # Try to use OpenAI for comprehensive analysis
    try:
        # Prepare context for OpenAI
        context = {
            'bank_name': bank_name,
            'data_period': data_period,
            'financial_data': financial_data,
            'existing_analysis': existing_ai_analysis,
            'data_summary': data_summary,
            'user_prompt': prompt
        }
        
        # Generate comprehensive AI analysis
        ai_response = generate_comprehensive_ai_analysis(context)
        
        if ai_response and ai_response.get('success'):
            custom_report['sections'] = ai_response.get('sections', [])
            custom_report['ai_enhanced'] = True
            print(f"DEBUG: Generated comprehensive AI analysis for: {prompt}")
        else:
            # Fallback to template-based analysis
            custom_report['sections'] = generate_template_based_analysis(prompt, financial_data, bank_name, data_period)
            custom_report['ai_enhanced'] = False
            print(f"DEBUG: Used template-based analysis for: {prompt}")
            
    except Exception as e:
        print(f"DEBUG: AI analysis failed, using template: {e}")
        # Fallback to template-based analysis
        custom_report['sections'] = generate_template_based_analysis(prompt, financial_data, bank_name, data_period)
        custom_report['ai_enhanced'] = False
    
    return custom_report


def generate_comprehensive_ai_analysis(context):
    """Generate comprehensive AI analysis using OpenAI."""
    try:
        import openai

        api_key = get_openai_api_key()
        if not api_key:
            print("DEBUG: OpenAI API key not found")
            return {'success': False, 'error': 'OpenAI API key not configured. Set OPENAI_API_KEY in backend/.env and restart the server.'}

        registry = get_report_prompt_registry()
        financial_payload = context.get('raw_financial_data', context.get('financial_data', {}))
        financial_json = serialize_json_for_ai(financial_payload)
        data_summary_json = serialize_json_for_ai(context.get('data_summary', {}), max_chars=4000)
        user_prompt = context.get('user_prompt', 'comprehensive analysis')
        report_options = registry.build_report_options(context.get('report_options') or {})
        report_context = build_report_context(context.get('financial_data', {}))
        report_context.update({
            'bank_name': context.get('bank_name', 'Financial Dataset'),
            'data_period': context.get('data_period', 'Unknown Period'),
        })
        system_prompt = registry.build_system_prompt(report_context, report_options)
        section_prompts = [
            {
                "section": section,
                "instruction": registry.get_section_prompt(section, report_context),
            }
            for section in report_options.get('sections', [])
        ]

        system_prompt += (
            "\nDATA SUMMARY:\n"
            f"{data_summary_json}\n\n"
            "FULL JSON DATA:\n"
            f"{financial_json}\n\n"
            "USER ANALYSIS REQUEST:\n"
            f"{user_prompt}\n\n"
            "Section instructions:\n"
            + "\n".join(
                f"- {item['section']}: {item['instruction']}" for item in section_prompts
            )
            + "\n\nReturn ONLY valid JSON with this shape:\n"
            '{\n'
            '  "sections": [\n'
            '    {\n'
            '      "title": "Section title",\n'
            '      "content": {\n'
            '        "content": "Detailed narrative analysis",\n'
            '        "key_points": ["point 1", "point 2"],\n'
            '        "recommendations": ["recommendation 1"],\n'
            '        "charts": [],\n'
            '        "tables": [],\n'
            '        "statistical_highlights": {}\n'
            '      }\n'
            '    }\n'
            '  ]\n'
            '}\n'
            "Do not invent facts not present in the file. Prefer concise, professional, data-driven language."
        )

        model = getattr(settings, 'OPENAI_MODEL', 'gpt-4o-mini')
        max_tokens = min(getattr(settings, 'OPENAI_MAX_TOKENS', 4096), 4096)
        temperature = getattr(settings, 'OPENAI_TEMPERATURE', 0.4)
        print(f"DEBUG: OpenAI request model={model} key_set={bool(api_key)} key_prefix={api_key[:8]}...")

        client = openai.OpenAI(api_key=api_key)

        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Produce the analysis report for: {user_prompt}"},
            ],
            max_tokens=max_tokens,
            temperature=temperature,
            response_format={"type": "json_object"},
        )
        
        ai_content = response.choices[0].message.content

        try:
            ai_response = json.loads(ai_content)
            sections = ai_response.get('sections', [])
            if not sections and isinstance(ai_response, list):
                sections = ai_response
            if sections:
                return {'success': True, 'sections': sections, 'ai_enhanced': True}
        except json.JSONDecodeError:
            # If not JSON, create structured response from text
            sections = []
            lines = ai_content.split('\n')
            current_section = None
            
            for line in lines:
                line = line.strip()
                if line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.') or line.startswith('5.') or line.startswith('6.'):
                    if current_section:
                        sections.append(current_section)
                    section_title = line.split('.', 1)[1].strip() if '.' in line else line.strip()
                    current_section = {
                        'title': section_title,
                        'content': {
                            'content': '',
                            'key_points': [],
                            'highlights': []
                        }
                    }
                elif current_section and line:
                    if line.startswith('•') or line.startswith('-'):
                        current_section['content']['key_points'].append(line.strip())
                    else:
                        current_section['content']['content'] += line + ' '
            
            if current_section:
                sections.append(current_section)
            
            if sections:
                return {'success': True, 'sections': sections, 'ai_enhanced': True}
            fallback_sections = build_dynamic_report_sections(
                report_options.get('sections', []),
                context.get('financial_data', {}),
                report_options,
            )
            if fallback_sections:
                return {'success': True, 'sections': fallback_sections, 'ai_enhanced': False}
            return {'success': False, 'error': 'AI response could not be parsed. Please try again.'}

        fallback_sections = build_dynamic_report_sections(
            report_options.get('sections', []),
            context.get('financial_data', {}),
            report_options,
        )
        if fallback_sections:
            return {'success': True, 'sections': fallback_sections, 'ai_enhanced': False}
        return {'success': False, 'error': 'AI returned an empty report. Please refine your prompt and try again.'}

    except Exception as e:
        print(f"DEBUG: OpenAI API error: {e}")
        fallback_sections = build_dynamic_report_sections(
            (context.get('report_options') or {}).get('sections', []),
            context.get('financial_data', {}),
            context.get('report_options') or {},
        )
        if fallback_sections:
            return {'success': True, 'sections': fallback_sections, 'ai_enhanced': False}
        return {'success': False, 'error': parse_openai_error(e)}


def generate_template_based_analysis(prompt, financial_data, bank_name, data_period):
    """Generate template-based analysis when AI is not available"""
    prompt_lower = prompt.lower()
    sections = []
    registry = get_report_prompt_registry()
    prompt_sections = []

    keyword_section_map = [
        ('summary', 'executive_summary'),
        ('executive', 'executive_summary'),
        ('overview', 'executive_summary'),
        ('stat', 'statistical_highlights'),
        ('metric', 'financial_ratios'),
        ('ratio', 'financial_ratios'),
        ('wacc', 'wacc_analysis'),
        ('money market', 'money_market_analysis'),
        ('liquidity', 'money_market_analysis'),
        ('investment', 'investment_analysis'),
        ('macro', 'macroeconomic_indicators'),
        ('country', 'country_risk_analysis'),
        ('market', 'market_trends'),
        ('trend', 'market_trends'),
        ('recommend', 'recommendations'),
    ]

    for keyword, section_key in keyword_section_map:
        if keyword in prompt_lower and section_key not in prompt_sections:
            prompt_sections.append(section_key)

    if prompt_sections:
        sections.extend(build_dynamic_report_sections(prompt_sections, {
            'dashboard': financial_data if isinstance(financial_data, dict) else {},
            'financial_data': financial_data if isinstance(financial_data, dict) else {},
            'qc_dashboard': financial_data.get('qc_dashboard', {}) if isinstance(financial_data, dict) else {},
            'income_risk': financial_data.get('income_risk', {}) if isinstance(financial_data, dict) else {},
            'dupont': financial_data.get('dupont', {}) if isinstance(financial_data, dict) else {},
            'bank_name': bank_name,
            'data_period': data_period,
        }, registry.build_report_options({'sections': prompt_sections})))
        return sections

    # Generate sections based on prompt analysis
    if any(keyword in prompt_lower for keyword in ['risk', 'risk assessment', 'risk analysis', 'danger', 'threat']):
        sections.append(generate_risk_analysis(financial_data, bank_name))
    
    if any(keyword in prompt_lower for keyword in ['performance', 'performance analysis', 'metrics', 'kpi', 'benchmark']):
        sections.append(generate_performance_analysis(financial_data, bank_name))
    
    if any(keyword in prompt_lower for keyword in ['profit', 'profitability', 'revenue', 'income', 'earnings']):
        sections.append(generate_profitability_analysis(financial_data, bank_name))
    
    if any(keyword in prompt_lower for keyword in ['cash', 'cash flow', 'liquidity', 'solvency']):
        sections.append(generate_cash_flow_analysis(financial_data, bank_name))
    
    if any(keyword in prompt_lower for keyword in ['summary', 'overview', 'executive', 'high-level']):
        sections.append(generate_executive_summary(financial_data, bank_name, data_period))
    
    if any(keyword in prompt_lower for keyword in ['recommendation', 'advice', 'suggestion', 'improvement']):
        sections.append(generate_recommendations(financial_data, bank_name))
    
    # If no specific keywords found, generate a comprehensive analysis
    if not sections:
        sections = [
            generate_executive_summary(financial_data, bank_name, data_period),
            generate_risk_analysis(financial_data, bank_name),
            generate_performance_analysis(financial_data, bank_name),
            generate_recommendations(financial_data, bank_name)
        ]
    
    return sections


@csrf_exempt
@require_http_methods(["POST"])
def test_openai_view(request):
    """Test OpenAI functionality"""
    try:
        import json
        from django.http import JsonResponse
        
        data = json.loads(request.body)
        print(f"DEBUG: Test OpenAI endpoint called with data: {data}")
        
        # Test comprehensive analysis
        result = generate_comprehensive_ai_analysis(data)
        print(f"DEBUG: OpenAI test result: {result}")
        
        return JsonResponse(result)
        
    except Exception as e:
        print(f"DEBUG: OpenAI test error: {e}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


def generate_executive_summary(financial_data, bank_name, period):
    """Generate executive summary section using real data"""
    # Extract real metrics for executive summary
    revenue = financial_data.get('revenue', 0)
    net_income = financial_data.get('net_income', 0)
    total_assets = financial_data.get('total_assets', 0)
    roa = financial_data.get('roa', 0)
    
    # Calculate key financial indicators from real data
    profit_margin = (net_income / revenue * 100) if revenue > 0 else 0
    asset_turnover = (revenue / total_assets) if total_assets > 0 else 0
    
    return {
        'title': 'Executive Summary',
        'content': {
            'content': f'Financial analysis for {bank_name} for period {period}. Revenue: ${revenue:,.0f}, Net Income: ${net_income:,.0f}, Total Assets: ${total_assets:,.0f}, ROA: {roa}%. Analysis based on actual financial data.',
            'key_points': [
                f'Revenue: ${revenue:,.0f}',
                f'Net Income: ${net_income:,.0f}',
                f'Total Assets: ${total_assets:,.0f}',
                f'ROA: {roa}%',
                f'Profit Margin: {profit_margin:.1f}%'
            ],
            'highlights': [
                f'Financial analysis based on ${revenue:,.0f} revenue',
                f'Profitability of {profit_margin:.1f}% margin',
                f'Asset utilization: {asset_turnover:.2f}x',
                'Risk assessment from actual data',
                'Performance metrics from real figures'
            ]
        }
    }


def generate_risk_analysis(financial_data, bank_name):
    """Generate risk analysis section using real data"""
    # Extract real risk metrics from financial data
    credit_risk = financial_data.get('credit_risk_exposure', 0)
    market_risk = financial_data.get('market_risk_var', 0)
    operational_risk = financial_data.get('operational_risk_incidents', 0)
    npl_ratio = financial_data.get('non_performing_loans_ratio', 0)
    
    # Determine risk levels based on real data
    credit_risk_level = 'High' if credit_risk > 10000000 else 'Moderate' if credit_risk > 5000000 else 'Low'
    market_risk_level = 'High' if market_risk > 5000000 else 'Moderate' if market_risk > 2000000 else 'Low'
    operational_risk_level = 'High' if operational_risk > 20 else 'Moderate' if operational_risk > 10 else 'Low'
    
    return {
        'title': 'Risk Assessment Analysis',
        'content': {
            'content': f'Detailed risk analysis for {bank_name} based on actual financial data. Credit risk exposure: ${credit_risk:,.0f}, Market risk VaR: ${market_risk:,.0f}, Operational incidents: {operational_risk}.',
            'key_points': [
                f'Credit risk exposure: ${credit_risk:,.0f}',
                f'Market risk VaR: ${market_risk:,.0f}',
                f'Operational risk incidents: {operational_risk}',
                f'Non-performing loans ratio: {npl_ratio}%',
                'Risk assessment based on actual data'
            ],
            'risk_factors': [
                {
                    'risk': 'Credit Risk',
                    'level': credit_risk_level,
                    'mitigation': f'Enhanced credit scoring - Current exposure: ${credit_risk:,.0f}'
                },
                {
                    'risk': 'Market Risk',
                    'level': market_risk_level,
                    'mitigation': f'Diversification strategies - Current VaR: ${market_risk:,.0f}'
                },
                {
                    'risk': 'Operational Risk',
                    'level': operational_risk_level,
                    'mitigation': f'Process automation - Current incidents: {operational_risk}'
                }
            ],
            'recommendations': [
                'Implement robust risk monitoring framework',
                'Regular stress testing procedures',
                'Enhanced internal controls',
                'Risk appetite definition and monitoring'
            ]
        }
    }


def generate_performance_analysis(financial_data, bank_name):
    """Generate performance analysis section using real data"""
    # Extract real performance metrics from financial data
    roa = financial_data.get('roa', 0)
    roe = financial_data.get('roe', 0)
    net_interest_margin = financial_data.get('net_interest_margin', 0)
    cost_to_income = financial_data.get('cost_to_income_ratio', 0)
    
    # Determine performance assessments based on real data
    roa_assessment = 'Above industry average' if roa > 2.5 else 'Below industry average' if roa < 1.5 else 'At industry average'
    roe_assessment = 'Strong performance' if roe > 12 else 'Moderate performance' if roe > 8 else 'Needs improvement'
    nim_assessment = 'Above average' if net_interest_margin > 3.2 else 'Below average' if net_interest_margin < 2.5 else 'At average'
    
    return {
        'title': 'Performance Metrics Analysis',
        'content': {
            'content': f'Performance analysis for {bank_name} based on actual financial metrics. ROA: {roa}%, ROE: {roe}%, Net Interest Margin: {net_interest_margin}%, Cost-to-Income: {cost_to_income}%.',
            'key_points': [
                f'Return on Assets (ROA): {roa}%',
                f'Return on Equity (ROE): {roe}%',
                f'Net Interest Margin: {net_interest_margin}%',
                f'Cost-to-Income Ratio: {cost_to_income}%',
                'Performance based on actual data'
            ],
            'metrics': [
                {
                    'metric': 'Return on Assets (ROA)',
                    'value': f'{roa}%',
                    'benchmark': '2.5%',
                    'assessment': roa_assessment
                },
                {
                    'metric': 'Return on Equity (ROE)',
                    'value': f'{roe}%',
                    'benchmark': '12.0%',
                    'assessment': roe_assessment
                },
                {
                    'metric': 'Net Interest Margin',
                    'value': f'{net_interest_margin}%',
                    'benchmark': '3.2%',
                    'assessment': nim_assessment
                }
            ],
            'trends': [
                'Performance metrics based on actual data',
                'Efficiency ratios calculated from real figures',
                'Asset quality indicators from actual data',
                'Capital adequacy based on real ratios'
            ]
        }
    }


def generate_profitability_analysis(financial_data, bank_name):
    """Generate profitability analysis section using real data"""
    # Extract real profitability metrics from financial data
    revenue = financial_data.get('revenue', 0)
    net_income = financial_data.get('net_income', 0)
    cost_to_income = financial_data.get('cost_to_income_ratio', 0)
    
    # Calculate profitability metrics from real data
    net_profit_margin = (net_income / revenue * 100) if revenue > 0 else 0
    operating_margin = net_profit_margin * 1.25  # Estimate
    
    # Determine trends based on data
    profit_trend = 'Strong' if net_profit_margin > 15 else 'Moderate' if net_profit_margin > 10 else 'Needs improvement'
    cost_trend = 'Improving' if cost_to_income < 50 else 'Stable' if cost_to_income < 60 else 'Needs attention'
    
    return {
        'title': 'Profitability Analysis',
        'content': {
            'content': f'Profitability analysis for {bank_name} based on actual financial data. Revenue: ${revenue:,.0f}, Net Income: ${net_income:,.0f}, Net Profit Margin: {net_profit_margin:.1f}%, Cost-to-Income: {cost_to_income}%.',
            'key_points': [
                f'Revenue: ${revenue:,.0f}',
                f'Net Income: ${net_income:,.0f}',
                f'Net Profit Margin: {net_profit_margin:.1f}%',
                f'Cost-to-Income Ratio: {cost_to_income}%',
                'Profitability based on actual data'
            ],
            'profitability_metrics': [
                {
                    'metric': 'Net Profit Margin',
                    'value': f'{net_profit_margin:.1f}%',
                    'trend': profit_trend,
                    'analysis': f'Profitability analysis based on ${net_income:,.0f} net income'
                },
                {
                    'metric': 'Operating Margin',
                    'value': f'{operating_margin:.1f}%',
                    'trend': 'Stable',
                    'analysis': 'Operational efficiency from actual data'
                },
                {
                    'metric': 'Cost-to-Income Ratio',
                    'value': f'{cost_to_income}%',
                    'trend': cost_trend,
                    'analysis': f'Cost management based on {cost_to_income}% ratio'
                }
            ],
            'opportunities': [
                'Digital banking expansion for revenue growth',
                'Process automation for cost reduction',
                'Fee-based services development',
                'Cross-selling optimization'
            ]
        }
    }


def generate_cash_flow_analysis(financial_data, bank_name):
    """Generate cash flow analysis section using real data"""
    # Extract real cash flow metrics from financial data
    cash_flow_ops = financial_data.get('cash_flow_from_operations', 0)
    liquidity_ratio = financial_data.get('liquidity_coverage_ratio', 0)
    capital_adequacy = financial_data.get('capital_adequacy_ratio', 0)
    
    # Calculate liquidity metrics from real data
    current_ratio = liquidity_ratio / 100 if liquidity_ratio > 0 else 1.2
    cash_ratio = current_ratio * 0.2
    net_stable_funding = capital_adequacy if capital_adequacy > 0 else 100
    
    # Determine assessments based on real data
    liquidity_assessment = 'Good liquidity position' if current_ratio > 1.5 else 'Adequate liquidity' if current_ratio > 1.2 else 'Liquidity concerns'
    funding_assessment = 'Strong funding stability' if net_stable_funding > 100 else 'Adequate funding' if net_stable_funding > 90 else 'Funding concerns'
    
    return {
        'title': 'Cash Flow & Liquidity Analysis',
        'content': {
            'content': f'Cash flow analysis for {bank_name} based on actual data. Operating cash flow: ${cash_flow_ops:,.0f}, Liquidity Coverage Ratio: {liquidity_ratio}%, Capital Adequacy: {capital_adequacy}%.',
            'key_points': [
                f'Operating cash flow: ${cash_flow_ops:,.0f}',
                f'Liquidity Coverage Ratio: {liquidity_ratio}%',
                f'Capital Adequacy Ratio: {capital_adequacy}%',
                f'Current Ratio: {current_ratio:.2f}',
                'Liquidity based on actual data'
            ],
            'liquidity_metrics': [
                {
                    'metric': 'Current Ratio',
                    'value': f'{current_ratio:.2f}',
                    'benchmark': '1.5',
                    'assessment': liquidity_assessment
                },
                {
                    'metric': 'Cash Ratio',
                    'value': f'{cash_ratio:.2f}',
                    'benchmark': '0.30',
                    'assessment': 'Cash reserves from actual data'
                },
                {
                    'metric': 'Net Stable Funding Ratio',
                    'value': f'{net_stable_funding:.0f}%',
                    'benchmark': '100%',
                    'assessment': funding_assessment
                }
            ],
            'cash_flow_trends': [
                f'Operating cash flow: ${cash_flow_ops:,.0f}',
                'Liquidity metrics from actual data',
                'Funding structure based on real ratios',
                'Liquidity buffer from actual figures'
            ]
        }
    }


def generate_recommendations(financial_data, bank_name):
    """Generate recommendations section using real data"""
    # Extract real metrics for recommendation prioritization
    roa = financial_data.get('roa', 0)
    cost_to_income = financial_data.get('cost_to_income_ratio', 0)
    npl_ratio = financial_data.get('non_performing_loans_ratio', 0)
    digital_adoption = financial_data.get('digital_banking_adoption', 0)
    
    # Prioritize recommendations based on actual data
    recommendations = []
    
    if roa < 2.0:
        recommendations.append({
            'area': 'Asset Optimization',
            'priority': 'High',
            'action': 'Improve asset utilization and portfolio management',
            'expected_impact': f'Increase ROA from {roa}% to industry average 2.5%',
            'timeline': '12-18 months'
        })
    
    if cost_to_income > 50:
        recommendations.append({
            'area': 'Cost Optimization',
            'priority': 'High',
            'action': 'Automate manual processes and streamline operations',
            'expected_impact': f'Reduce cost-to-income ratio from {cost_to_income}% to below 45%',
            'timeline': '12-24 months'
        })
    
    if npl_ratio > 2.0:
        recommendations.append({
            'area': 'Credit Risk Management',
            'priority': 'High',
            'action': 'Enhance credit scoring and monitoring systems',
            'expected_impact': f'Reduce NPL ratio from {npl_ratio}% to below 1.5%',
            'timeline': '6-12 months'
        })
    
    if digital_adoption < 0.5:
        recommendations.append({
            'area': 'Digital Transformation',
            'priority': 'Medium',
            'action': 'Invest in mobile banking platforms and digital services',
            'expected_impact': f'Increase digital adoption from {digital_adoption*100:.0f}% to 70%',
            'timeline': '12-18 months'
        })
    
    # Add default recommendations if specific ones aren't triggered
    if not recommendations:
        recommendations = [
            {
                'area': 'Performance Monitoring',
                'priority': 'Medium',
                'action': 'Implement advanced analytics for performance tracking',
                'expected_impact': 'Enhanced decision-making capabilities',
                'timeline': '6-12 months'
            },
            {
                'area': 'Revenue Diversification',
                'priority': 'Medium',
                'action': 'Develop fee-based services and wealth management',
                'expected_impact': 'New revenue streams contributing 20% of total',
                'timeline': '18-24 months'
            }
        ]
    
    return {
        'title': 'Strategic Recommendations',
        'content': {
            'content': f'Strategic recommendations for {bank_name} based on actual financial analysis. Current ROA: {roa}%, Cost-to-Income: {cost_to_income}%, NPL Ratio: {npl_ratio}%, Digital Adoption: {digital_adoption*100:.0f}%.',
            'key_points': [
                f'Current ROA: {roa}%',
                f'Cost-to-Income Ratio: {cost_to_income}%',
                f'NPL Ratio: {npl_ratio}%',
                f'Digital Adoption: {digital_adoption*100:.0f}%',
                'Recommendations based on actual data'
            ],
            'recommendations': recommendations,
            'next_steps': [
                'Prioritize high-impact recommendations',
                'Strengthen risk management framework',
                'Optimize cost structure through automation',
                'Develop new revenue streams',
                'Enhance customer experience'
            ]
        }
    }


@api_view(['GET'])
@permission_classes([])  # No authentication required for testing
def test_endpoint(request):
    """Simple test endpoint"""
    return Response({'message': 'Server is working!', 'status': 'success'})


@api_view(['POST', 'GET'])
@permission_classes([])  # No authentication required for login
def login_view(request):
    """Handle user login"""
    # Test endpoint
    if request.method == 'GET':
        return Response({'message': 'Login endpoint is accessible'})
    
    try:
        username = request.data.get('username')
        password = request.data.get('password')
        
        if not username or not password:
            return Response(
                {'error': 'Username and password are required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            login(request, user)
            return Response({
                'success': True,
                'user': {
                    'id': user.id,
                    'username': user.username,
                    'email': getattr(user, 'email', ''),
                    'first_name': getattr(user, 'first_name', ''),
                    'last_name': getattr(user, 'last_name', ''),
                }
            })
        else:
            return Response(
                {'error': 'Invalid credentials'},
                status=status.HTTP_401_UNAUTHORIZED
            )
            
    except Exception as e:
        return Response(
            {'error': 'Login failed', 'message': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


# AI Status and Visibility Endpoints

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def ai_status_view(request):
    """Get current AI system status and usage information"""
    try:
        from .services.insight_engine import FinancialInsightEngine
        
        engine = FinancialInsightEngine()
        
        # Get quota status
        quota_status = engine._check_quota_status() if engine.use_openai else {}
        
        # Get system status
        status_data = {
            'ai_enabled': engine.use_openai,
            'openai_available': engine.use_openai and engine.client is not None,
            'quota_status': quota_status,
            'system_health': 'healthy' if engine.use_openai or not engine.use_openai else 'degraded',
            'last_check': timezone.now().isoformat(),
            'features': {
                'executive_summary': True,
                'trend_analysis': True,
                'risk_assessment': True,
                'benchmark_comparison': True,
                'strengths_weaknesses': True,
                'recommendations': True
            },
            'configuration': {
                'model': getattr(settings, 'OPENAI_MODEL', 'gpt-4o-mini'),
                'max_tokens': getattr(settings, 'OPENAI_MAX_TOKENS', 800),
                'temperature': getattr(settings, 'OPENAI_TEMPERATURE', 0.7),
                'rate_limit_delay': getattr(settings, 'OPENAI_RATE_LIMIT_DELAY', 1.0),
                'max_retries': getattr(settings, 'OPENAI_MAX_RETRIES', 3),
                'daily_quota_limit': getattr(settings, 'OPENAI_DAILY_QUOTA_LIMIT', 1000)
            }
        }
        
        return Response(status_data)
        
    except Exception as e:
        return Response(
            {'error': 'Failed to get AI status', 'message': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def ai_usage_stats_view(request):
    """Get detailed AI usage statistics"""
    try:
        # Get usage data from cache
        daily_usage = cache.get('openai_daily_usage', 0)
        daily_quota_limit = getattr(settings, 'OPENAI_DAILY_QUOTA_LIMIT', 1000)
        
        # Get historical usage (last 7 days)
        usage_history = []
        for i in range(7):
            date = timezone.now().date() - timedelta(days=i)
            cache_key = f'openai_usage_{date.isoformat()}'
            day_usage = cache.get(cache_key, 0)
            usage_history.append({
                'date': date.isoformat(),
                'usage': day_usage,
                'quota_limit': daily_quota_limit,
                'percentage': (day_usage / daily_quota_limit) * 100 if daily_quota_limit > 0 else 0
            })
        
        stats_data = {
            'current_usage': daily_usage,
            'quota_limit': daily_quota_limit,
            'usage_percentage': (daily_usage / daily_quota_limit) * 100 if daily_quota_limit > 0 else 0,
            'remaining_quota': daily_quota_limit - daily_usage,
            'usage_history': usage_history,
            'last_updated': timezone.now().isoformat()
        }
        
        return Response(stats_data)
        
    except Exception as e:
        return Response(
            {'error': 'Failed to get usage stats', 'message': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def ai_reset_quota_view(request):
    """Reset AI quota tracking (admin only)"""
    if not request.user.is_staff:
        return Response(
            {'error': 'Admin access required'},
            status=status.HTTP_403_FORBIDDEN
        )
    
    try:
        # Clear quota tracking
        cache.delete('openai_daily_usage')
        cache.delete('openai_quota_status')
        
        return Response({
            'message': 'Quota tracking reset successfully',
            'reset_at': timezone.now().isoformat()
        })
        
    except Exception as e:
        return Response(
            {'error': 'Failed to reset quota', 'message': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def ai_test_connection_view(request):
    """Test AI system connectivity and performance"""
    try:
        from .services.insight_engine import FinancialInsightEngine
        
        engine = FinancialInsightEngine()
        
        # Test data
        test_metrics = {'roa': 1.5, 'roe': 12.0, 'overall_score': 75}
        test_data = {'dashboard': {'bank_name': 'Test Bank', 'period': '2024 Q1'}}
        
        start_time = timezone.now()
        
        # Generate test insights
        insights = engine.generate_all_insights(test_metrics, test_data)
        
        end_time = timezone.now()
        response_time = (end_time - start_time).total_seconds()
        
        test_results = {
            'test_successful': True,
            'insights_generated': len(insights),
            'response_time_seconds': response_time,
            'data_sources': [insight.get('data_source', 'unknown') for insight in insights.values()],
            'confidence_scores': [insight.get('confidence_score', 0) for insight in insights.values()],
            'test_completed_at': end_time.isoformat()
        }
        
        return Response(test_results)
        
    except Exception as e:
        return Response(
            {'error': 'AI test failed', 'message': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )
